from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import os

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        # Use all-MiniLM-L6-v2 model for embeddings
        self.embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        self.vector_store = None
        self.full_text = []

    def load_document(self, file_path):
        """Load and process a Word document."""
        try:
            if not os.path.exists(file_path):
                print(f"Document not found at path: {file_path}")
                return False

            print(f"Loading document from: {file_path}")
            doc = Document(file_path)
            self.full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    self.full_text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            self.full_text.append(cell.text.strip())
            
            print(f"Extracted {len(self.full_text)} text segments from document")
            
            if not self.full_text:
                print("Warning: No text content found in document")
                return False
            
            # Split text into chunks
            texts = self.text_splitter.create_documents(self.full_text)
            print(f"Created {len(texts)} text chunks")
            
            # Create vector store
            self.vector_store = FAISS.from_documents(texts, self.embeddings)
            print("Vector store created successfully")
            
            return True
        
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return False

    def search_similar_content(self, query, k=3):
        """Search for similar content in the vector store."""
        if not self.vector_store:
            print("Warning: Vector store not initialized")
            return []
        
        try:
            results = self.vector_store.similarity_search(query, k=k)
            print(f"Found {len(results)} similar content segments")
            return [doc.page_content for doc in results]
        except Exception as e:
            print(f"Error searching content: {str(e)}")
            return []

    def get_document_topics(self):
        """Get main topics from the document."""
        if not self.full_text:
            print("Warning: No document content available")
            return []
        
        print(f"Returning {len(self.full_text)} text segments")
        return self.full_text 