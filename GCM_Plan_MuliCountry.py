import pandas as pd
from docx import Document
import datetime
import requests
import time
import logging
from collections import defaultdict
import docx

class LLMRequest:
    def __init__(self, auth_token, user_id):
        self.auth_token = auth_token
        self.user_id = user_id
        self.base_url = "https://prod.ulchatbot.com/api/chats/chat/advanced"
        self.file_path = r"C:\Users\104288\.cursor-tutor\Cursor_GCM\Compare_ByCountry.xlsx"
        self.data_df = None
        self.countries = []
        self.product_name = ""
        self.rate_limit = 1.0  # 1 second between requests
        self.last_request_time = 0
        self.load_excel_data()

    def load_excel_data(self):
        """Load data from the Data tab of the Excel file."""
        print("Loading data from:", self.file_path)
        xls = pd.ExcelFile(self.file_path)
        
        # Load regulatory data from the Data sheet
        self.data_df = pd.read_excel(xls, sheet_name='Data')
        
        # Get country columns (all columns after 'Attribute_Name')
        attribute_col_idx = self.data_df.columns.get_loc('Attribute_Name')
        self.countries = self.data_df.columns[attribute_col_idx + 1:].tolist()
        print(f"Countries extracted for comparison: {self.countries}")

        # Load product information from the Product sheet
        product_df = pd.read_excel(xls, sheet_name='Product')
        self.product_name = product_df.iloc[0]['Product_Category']
        print(f"Product loaded: {self.product_name}")

    def call_api(self, messages):
        """Call the LLM API with rate limiting and retries"""
        current_time = time.time()
        time_since_last_request = current_time - self.last_request_time
        if time_since_last_request < self.rate_limit:
            time.sleep(self.rate_limit - time_since_last_request)

        headers = {
            "accept": "application/json",
            "accept-encoding": "gzip, deflate, br, zstd",
            "accept-language": "en-US,en;q=0.9",
            "authorization": f"Bearer {self.auth_token}",
            "content-type": "application/json",
            "cookie": "tosAccepted=true",
            "origin": "https://prod.ulchatbot.com",
            "referer": "https://prod.ulchatbot.com/",
            "sec-ch-ua": '"Microsoft Edge";v="131", "Chromium";v="131", "Not_A Brand";v="24"',
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": '"Windows"',
            "sec-fetch-dest": "empty",
            "sec-fetch-mode": "cors",
            "sec-fetch-site": "same-origin",
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0"
        }
        
        payload = {
            "approach": "RetrieveThenRead",
            "userId": self.user_id,
            "conversationId": "d6929919-701d-4d12-7b57-08dd1f597649",
            "history": messages,
            "isRagUsed": False
        }

        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = requests.post(self.base_url, headers=headers, json=payload)
                self.last_request_time = time.time()
                
                if response.status_code != 200:
                    logging.error(f"Response text: {response.text}")
                response.raise_for_status()
                return response.json()
            except requests.exceptions.RequestException as e:
                if attempt == max_retries - 1:
                    logging.error(f"API request failed after {max_retries} attempts: {str(e)}")
                    raise
                time.sleep(1 * (attempt + 1))

    def analyze_regional_patterns(self):
        """Analyze patterns in requirements across countries to identify regional commonalities"""
        regional_patterns = defaultdict(lambda: defaultdict(list))
        
        for category in self.data_df['Regulation_Category'].unique():
            category_data = self.data_df[self.data_df['Regulation_Category'] == category]
            
            for _, row in category_data.iterrows():
                # Group countries with identical requirements
                requirements = defaultdict(list)
                for country in self.countries:
                    req = str(row[country]).strip()
                    requirements[req].append(country)
                
                # Store patterns
                for req, countries in requirements.items():
                    if len(countries) > 1:  # Only store if shared by multiple countries
                        regional_patterns[category][row['Attribute_Name']].append({
                            'requirement': req,
                            'countries': countries
                        })
        
        return regional_patterns

    def generate_country_action_summary(self):
        """Generate a summary of required actions for each country"""
        country_actions = {country: defaultdict(list) for country in self.countries}
        shared_actions = defaultdict(list)
        
        for _, row in self.data_df.iterrows():
            category = row['Regulation_Category']
            attribute = row['Attribute_Name']
            
            # Group countries with identical requirements
            requirements = defaultdict(list)
            for country in self.countries:
                req = str(row[country]).strip()
                if req and req.lower() != 'nan':
                    requirements[req].append(country)
            
            # Process requirements
            for req, req_countries in requirements.items():
                if len(req_countries) > 1:
                    # Shared requirement - add to first country and mark as leverage for others
                    primary_country = req_countries[0]
                    action_item = f"{category} - {attribute}: {req}"
                    country_actions[primary_country][category].append(f"{action_item}")
                    for other_country in req_countries[1:]:
                        country_actions[other_country][category].append(
                            f"Leverage {primary_country}'s compliance for: {attribute}"
                        )
                else:
                    # Unique requirement
                    country = req_countries[0]
                    country_actions[country][category].append(
                        f"{category} - {attribute}: {req}"
                    )
        
        return country_actions

    def add_action_summary_table(self, doc, country_actions):
        """Add a formatted table summarizing required actions by country"""
        doc.add_heading('Country-Specific Action Summary', level=1)
        doc.add_paragraph('This section provides a comprehensive checklist of required actions for each country, including opportunities for leveraging common requirements.')
        
        for country in self.countries:
            doc.add_heading(f'{country} Action Items', level=2)
            actions = country_actions[country]
            
            if not any(actions.values()):
                doc.add_paragraph('No specific actions required.')
                continue
                
            for category in self.data_df['Regulation_Category'].unique():
                if category in actions and actions[category]:
                    doc.add_heading(category, level=3)
                    action_list = doc.add_paragraph()
                    for action in actions[category]:
                        action_list.add_run(f"• {action}\n")

    def generate_compliance_plan(self):
        """Generate a comprehensive compliance plan document with detailed attribute analysis"""
        print("Generating compliance plan document...")
        doc = Document()
        
        # Title and Introduction
        title = doc.add_heading(f'Multi-Country Regulatory Compliance Strategy for {self.product_name}', 0)
        title.alignment = 1  # Center alignment
        doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        # Add style for better formatting
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = docx.shared.Pt(11)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_prompt = f"""Create a concise executive summary for the regulatory compliance strategy for {self.product_name} targeting: {', '.join(self.countries)}.
Focus on:
1. Overview of regulatory scope
2. Key compliance requirements
3. Critical considerations for market entry

Please use clear, professional language without special characters or symbols."""
        
        summary_response = self.call_api([{
            "role": "user",
            "content": summary_prompt,
            "date": datetime.datetime.utcnow().isoformat() + "Z"
        }])
        doc.add_paragraph(summary_response['choices'][0]['message']['content'])
        
        # Regulatory Categories Overview
        doc.add_heading('Regulatory Categories Overview', level=1)
        categories = self.data_df['Regulation_Category'].unique()
        overview_text = "This compliance plan covers the following regulatory categories:\n\n"
        for category in categories:
            overview_text += f"{category}\n"
        doc.add_paragraph(overview_text)
        
        # Detailed Requirements Analysis by Category and Attribute
        doc.add_heading('Detailed Requirements Analysis', level=1)
        
        for category in categories:
            doc.add_heading(f'{category}', level=2)
            category_data = self.data_df[self.data_df['Regulation_Category'] == category]
            
            # Analyze each attribute within the category
            for _, row in category_data.iterrows():
                attribute_name = row['Attribute_Name']
                doc.add_heading(attribute_name, level=3)
                
                # Create a detailed analysis prompt for this specific attribute
                attribute_prompt = f"""Analyze the regulatory requirements for {self.product_name} regarding {attribute_name} across {', '.join(self.countries)}.

Raw requirement data:
{dict(zip(self.countries, [row[country] for country in self.countries]))}

Please provide:
1. Common Requirements: List the exact requirements that are identical across countries, including specific standards, dates, and values mentioned.
2. Country-Specific Requirements: For each country with unique requirements, specify the exact requirements as stated in the data.
3. Implementation Considerations: Key points for compliance with these specific requirements.

Format the response in clear paragraphs without special characters or symbols. Use professional business language."""
                
                attribute_response = self.call_api([{
                    "role": "user",
                    "content": attribute_prompt,
                    "date": datetime.datetime.utcnow().isoformat() + "Z"
                }])
                doc.add_paragraph(attribute_response['choices'][0]['message']['content'])
        
        # Implementation Strategy
        doc.add_heading('Implementation Strategy', level=1)
        strategy_prompt = f"""Based on the detailed regulatory requirements analyzed above, create an implementation strategy for {self.product_name} across {len(self.countries)} markets.

Consider:
1. Sequence of compliance activities
2. Resource requirements
3. Timeline considerations
4. Risk mitigation approaches

Focus on practical steps and specific requirements identified in the analysis. Use clear business language without special characters."""
        
        strategy_response = self.call_api([{
            "role": "user",
            "content": strategy_prompt,
            "date": datetime.datetime.utcnow().isoformat() + "Z"
        }])
        doc.add_paragraph(strategy_response['choices'][0]['message']['content'])
        
        # Add Country-Specific Action Summary
        country_actions = self.generate_country_action_summary()
        self.add_action_summary_table(doc, country_actions)
        
        # Save document with improved formatting
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                paragraph.space_before = docx.shared.Pt(12)
                paragraph.space_after = docx.shared.Pt(6)
            else:
                paragraph.space_after = docx.shared.Pt(10)
        
        doc_path = r"C:\Users\104288\.cursor-tutor\Cursor_GCM\Multi_Country_Compliance_Plan.docx"
        doc.save(doc_path)
        print(f"Compliance plan saved to: {doc_path}")

if __name__ == "__main__":
    auth_token = 'eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiIsIng1dCI6InoxcnNZSEhKOS04bWdndDRIc1p1OEJLa0JQdyIsImtpZCI6InoxcnNZSEhKOS04bWdndDRIc1p1OEJLa0JQdyJ9.eyJhdWQiOiJhcGk6Ly83NzMxZWM3NC1mMWNhLTQxZDUtYTBlNS04MWY2ODM3YWMyNDciLCJpc3MiOiJodHRwczovL3N0cy53aW5kb3dzLm5ldC83MDExNTk1NC0wY2NkLTQ1ZjAtODdiZC0wM2IyYTM1ODc1NjkvIiwiaWF0IjoxNzM2NTM4MDA4LCJuYmYiOjE3MzY1MzgwMDgsImV4cCI6MTczNjU0MjUxNSwiYWNyIjoiMSIsImFpbyI6IkFWUUFxLzhZQUFBQTB3Q05FUDh0UEloOEMxMnJRbGhlbmhUeHV1MEFRNTU2K2JxQTkxbDBTNHNtSVN6WklJZXk2V1dkSTNHVzNXSkt5NHBSZldwMFJBMExyempsVS9QVWFLZzc3T0FmVGJYY1BEU0tVWE54ZkNvPSIsImFtciI6WyJwd2QiLCJyc2EiLCJtZmEiXSwiYXBwaWQiOiI3NzMxZWM3NC1mMWNhLTQxZDUtYTBlNS04MWY2ODM3YWMyNDciLCJhcHBpZGFjciI6IjAiLCJkZXZpY2VpZCI6IjMyOTk1Y2U1LTM0NjMtNGE4Yy1iM2MyLWFkNjRlZmY0ZjhkNyIsImZhbWlseV9uYW1lIjoiWWFuZyIsImdpdmVuX25hbWUiOiJKZW5ueSIsImlwYWRkciI6IjIwNS4xNzguMzAuNDUiLCJuYW1lIjoiWWFuZywgSmVubnkiLCJvaWQiOiI1YTVlNDNjNi0wMWEzLTQ4MjItOGE0My0yY2E2OThlYjViNTYiLCJvbnByZW1fc2lkIjoiUy0xLTUtMjEtMzYzODA4OTg2OC0zMDgxMTc1MTE1LTMxMzI5NzcyOS0zMjE1MzMiLCJyaCI6IjEuQVJzQVZGa1JjTTBNOEVXSHZRT3lvMWgxYVhUc01YZks4ZFZCb09XQjlvTjZ3a2NiQU9rYkFBLiIsInNjcCI6IkFwaS5SZWFkIiwic3ViIjoiRTFWeklZblc2RWI0aE5ZU0xZNE9LUTVBbXQ4SVBUaVRPWXZCSHZaNzYxZyIsInRpZCI6IjcwMTE1OTU0LTBjY2QtNDVmMC04N2JkLTAzYjJhMzU4NzU2OSIsInVuaXF1ZV9uYW1lIjoiMTA0Mjg4QGdsb2JhbC51bC5jb20iLCJ1cG4iOiIxMDQyODhAZ2xvYmFsLnVsLmNvbSIsInV0aSI6IndxTWFhaF9ndEVxNHNoMFVFQ1EtQUEiLCJ2ZXIiOiIxLjAifQ.eb8rDFkRtheAj_5PUoTjVVheDYXQuCBqGE0fQ9TbetFKWwfYfHCq09BufkUar6bm1im383Ldc_9Y2mngB69AvYflDQnJpw-FZmDxMWBPU7etisCO1P9j9SlIV18lm-jHMaOy23e33Deq70aZ5op3ecotRUANB2pTkrG8M8kG9dQSd1_m02iRFYhhCItJwAtVz6KFZeEdYG3xG9126dLBYOYJu3mQQ32DCUKPDZc4pQ9k-Hetb40OgGA-z1TMGBs_hySR7OYNgLnmx-7bjQQruEXKJj8D5PHLCXWhix6JqGB9CJdOLej5V53kRouNYpGNpIGcjo8Ro48fFvuWTkP7ag'
    user_id = '5a5e43c6-01a3-4822-8a43-2ca698eb5b56'
    
    llm_request = LLMRequest(auth_token, user_id)
    llm_request.generate_compliance_plan()
