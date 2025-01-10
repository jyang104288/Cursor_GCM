import pandas as pd
import os
from openpyxl import Workbook
from docx import Document
import logging
import time
import requests
import datetime

# change 1. 


# Configure logging
logging.basicConfig(filename='debug.log', level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')


class LLMRequest:
    def __init__(self, auth_token, user_id):
        self.auth_token = auth_token
        self.user_id = user_id
        self.base_url = "https://prod.ulchatbot.com/api/chats/chat/advanced"
        self.file_path = r"C:\Users\104288\.cursor-tutor\CompareReport\Compare_ByCountry.xlsx"
        self.rate_limit = 0.5  # Minimum time between requests in seconds
        self.last_request_time = 0
        self.search_domains = [
            "ul.com",
            "gov",
            "standards.gov",
            "iso.org",
            "iec.ch",
            "astm.org",
            "nist.gov",
            "cpsc.gov"
        ]
        self.load_excel_data()
        self.rate_limit = 1.0  # Increase to 1 second between requests
        print("LLMRequest initialized with user_id:", user_id)

    def compare_regulatory_info(self):
        """Process regulatory queries and generate both Excel and Word outputs"""
        try:
            regulatory_data = self.data_df[['Regulation_Category', 'Regulation_Subcategory', 
                                         'Attribute_Name', self.country1, self.country2]]
            summary_results = []

            for index, row in regulatory_data.iterrows():
                print(f"Comparing attributes for row {index}: {row[self.country1]} vs {row[self.country2]}")
                summary = self.compare_attributes(row)  # Call the compare_attributes method
                summary_results.append({
                    'Regulation_Category': row['Regulation_Category'],
                    'Regulation_Subcategory': row['Regulation_Subcategory'],
                    'Attribute_Name': row['Attribute_Name'],
                    self.country1: row[self.country1],
                    self.country2: row[self.country2],
                    'Summary': summary
                })
                print(f"Summary for row {index}: {summary}")

            # Convert summary results to a DataFrame
            output_df = pd.DataFrame(summary_results)

            # Generate the compliance plan document
            self.generate_compliance_plan(output_df)  # Call to generate the compliance plan

        except Exception as e:
            print(f"Error during comparison: {e}")

    def load_excel_data(self):
        """Load data from Excel file"""
        print("Loading Excel data from:", self.file_path)
        xls = pd.ExcelFile(self.file_path)
        
        # Load country names from the Compare sheet
        compare_df = pd.read_excel(xls, sheet_name='Compare')
        
        # Accessing country names from the second row (index 1)
        self.country1 = compare_df.iloc[0, 0]  # Cell A2
        self.country2 = compare_df.iloc[0, 1]  # Cell B2
        print(f"Selected countries for comparison: {self.country1}, {self.country2}")

        # Load regulatory data from the Data sheet
        self.data_df = pd.read_excel(xls, sheet_name='Data')
        print(f"Loaded regulatory data with {len(self.data_df)} rows")

    def call_api(self, messages):
        print("Calling API with messages:", messages)
        # Rate limiting
        current_time = time.time()
        time_since_last_request = current_time - self.last_request_time
        if time_since_last_request < self.rate_limit:
            time.sleep(self.rate_limit - time_since_last_request)

        headers = {
            "accept": "application/json",
            "accept-encoding": "gzip, deflate, br, zstd",
            "accept-language": "en-US,en;q=0.9,zh;q=0.8",
            "authorization": f"Bearer {self.auth_token}",
            "content-type": "application/json",
            "cookie": "tosAccepted=true",
            "origin": "https://prod.ulchatbot.com",
            "referer": "https://prod.ulchatbot.com/",
            "sec-ch-ua": '"Microsoft Edge";v="131", "Chromium";v="131", "Not_A Brand";v="24"',
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": '"Windows"',
            "sec-fetch-dest": "empty",
            "sec-fetch-mode": "cors",
            "sec-fetch-site": "same-origin",
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0"
        }
        
        payload = {
            "approach": "RetrieveThenRead",
            "userId": self.user_id,
            "conversationId": "feabf9d6-7249-45ba-7b1d-08dd1f597649",
            "history": messages,
            "isRagUsed": False
        }

        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"Attempt {attempt + 1} to call API")
                response = requests.post(self.base_url, headers=headers, json=payload)
                self.last_request_time = time.time()
                
                if response.status_code != 200:
                    logging.error(f"Response text: {response.text}")
                response.raise_for_status()
                print("API response received:", response.json())
                return response.json()
            except requests.exceptions.RequestException as e:
                print(f"API request failed: {str(e)}")
                if attempt == max_retries - 1:
                    logging.error(f"API request failed after {max_retries} attempts: {str(e)}")
                    raise
                time.sleep(1 * (attempt + 1))  # Exponential backoff

    def compare_attributes(self, row):
        """Compare two attributes using LLM and return a summary"""
        try:
            # Extract values from the row
            attr1 = str(row[self.country1]).strip()
            attr2 = str(row[self.country2]).strip()
            category = row['Regulation_Category']
            attribute = row['Attribute_Name']

            # Modified prompt to handle multiple information comparison
            final_prompt = (
                f"Compare these two sets of information:\n"
                f"Set 1: {attr1}\n"
                f"Set 2: {attr2}\n\n"
                "Rules for comparison:\n"
                "1. If both sets are completely identical, respond with: 'same'\n"
                "2. If there are differences, respond in this format:\n"
                "Common Elements: [list items that appear in both sets]\n"
                f"Only in Set 1: [list items unique to Set 1]\n"
                f"Only in Set 2: [list items unique to Set 2]\n"
                "3. Treat comma-separated items as separate elements\n"
                "4. Compare exact content including spacing and punctuation\n"
                "5. Do not provide any additional analysis"
            )

            messages = [{
                "role": "user",
                "content": final_prompt,
                "date": datetime.datetime.utcnow().isoformat() + "Z"
            }]

            response = self.call_api(messages)
            ai_answer = response['choices'][0]['message']['content'].strip()
            
            return ai_answer
         # Format the final response with context
            formatted_response = (
                f"Regulatory Category: {category}\n"
                f"Attribute: {attribute}\n"
                f"{ai_answer.replace('Set 1', self.country1).replace('Set 2', self.country2)}"
            )

            return formatted_response

        except Exception as e:
            print(f"Error comparing attributes: {e}")
            return None

    def generate_compliance_plan(self, regulatory_data):
        """Generate a Word document for the compliance plan based on regulatory requirements using LLM"""
        print("Generating compliance plan document...")
        
        doc = Document()
        
        # Add title and header
        doc.add_heading(f'Regulatory Compliance Strategy for Cooktop Products - {self.country1} and {self.country2}', 0)
        doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        # Regulatory Scope Summary
        doc.add_heading('Regulatory Scope Summary', level=1)
        unique_categories = regulatory_data['Regulation_Category'].unique()
        doc.add_paragraph("This compliance plan covers the following regulatory categories:")
        for category in unique_categories:
            doc.add_paragraph(f"• {category}", style='List Bullet')
        doc.add_paragraph("\nNote: This compliance plan is limited to the regulatory categories listed above based on the provided data. Additional regulatory requirements may apply.")
        
        # Project Overview
        doc.add_heading('Project Overview', level=1)
        overview_prompt = f"""As a product compliance manager, provide a concise project overview for a cooktop product targeting {self.country1} and {self.country2} markets.
        Focus on shared requirements and key differences based on this regulatory data:
        {regulatory_data[['Regulation_Category', self.country1, self.country2]].to_string()}"""
        
        overview_response = self.call_api([{
            "role": "user",
            "content": overview_prompt
        }])
        doc.add_paragraph(overview_response['choices'][0]['message']['content'])
        
        # Certification Requirements
        doc.add_heading('Certification Requirements', level=1)
        cert_prompt = f"""Based on this regulatory data for {self.country1} and {self.country2}:
        {regulatory_data[regulatory_data['Regulation_Category'].str.contains('Certification|Standard', case=False)].to_string()}
        
        List the following:
        1. Primary certifications required
        2. Most efficient certification path
        3. Validity periods
        4. Any country-specific requirements"""
        
        cert_response = self.call_api([{
            "role": "user",
            "content": cert_prompt
        }])
        doc.add_paragraph(cert_response['choices'][0]['message']['content'])
        
        # Testing Strategy
        doc.add_heading('Testing Strategy', level=1)
        testing_prompt = f"""Based on the standards and requirements for {self.country1} and {self.country2}:
        {regulatory_data[regulatory_data['Regulation_Category'].str.contains('Test|Standard|Safety', case=False)].to_string()}
        
        Create a consolidated testing plan including:
        1. Required safety testing
        2. EMC testing requirements
        3. Testing optimization strategies
        4. Laboratory requirements
        5. Test report acceptability"""
        
        testing_response = self.call_api([{
            "role": "user",
            "content": testing_prompt
        }])
        doc.add_paragraph(testing_response['choices'][0]['message']['content'])
        
        # Documentation Requirements
        doc.add_heading('Documentation Requirements', level=1)
        doc_prompt = f"""Based on documentation requirements for {self.country1} and {self.country2}:
        {regulatory_data[regulatory_data['Regulation_Category'].str.contains('Document|Manual|Label', case=False)].to_string()}
        
        Detail the documentation requirements for:
        1. Technical File contents
        2. Compliance documentation
        3. User documentation and language requirements
        4. Labeling requirements"""
        
        doc_response = self.call_api([{
            "role": "user",
            "content": doc_prompt
        }])
        doc.add_paragraph(doc_response['choices'][0]['message']['content'])
        
        # Implementation Timeline
        doc.add_heading('Implementation Timeline', level=1)
        timeline_prompt = f"""Create a detailed implementation timeline for achieving compliance in {self.country1} and {self.country2}, including:
        1. Pre-certification phase
        2. Certification phase
        3. Production phase
        4. Estimated duration for each phase
        5. Key milestones and dependencies"""
        
        timeline_response = self.call_api([{
            "role": "user",
            "content": timeline_prompt
        }])
        doc.add_paragraph(timeline_response['choices'][0]['message']['content'])
        
        # Cost Optimization Strategy
        doc.add_heading('Cost Optimization Strategy', level=1)
        cost_prompt = f"""Suggest cost optimization strategies for compliance in {self.country1} and {self.country2}, considering:
        1. Certification efficiency
        2. Testing efficiency
        3. Documentation efficiency
        4. Resource optimization"""
        
        cost_response = self.call_api([{
            "role": "user",
            "content": cost_prompt
        }])
        doc.add_paragraph(cost_response['choices'][0]['message']['content'])
        
        # Risk Mitigation
        doc.add_heading('Risk Mitigation', level=1)
        risk_prompt = f"""Outline a risk mitigation strategy for maintaining compliance in {self.country1} and {self.country2}, including:
        1. Certificate validity management
        2. Compliance monitoring
        3. Documentation control
        4. Quality system maintenance"""
        
        risk_response = self.call_api([{
            "role": "user",
            "content": risk_prompt
        }])
        doc.add_paragraph(risk_response['choices'][0]['message']['content'])
        
        # Save document
        compliance_doc_path = os.path.join(r"C:\Users\104288\.cursor-tutor\CompareReport", 
                                            f"Compliance_Plan_{self.country1}_{self.country2}.docx")
        doc.save(compliance_doc_path)
        print(f"Compliance plan document saved to: {compliance_doc_path}")

if __name__ == "__main__":
    # Define your actual auth_token and user_id here
    auth_token = 'eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiIsIng1dCI6InoxcnNZSEhKOS04bWdndDRIc1p1OEJLa0JQdyIsImtpZCI6InoxcnNZSEhKOS04bWdndDRIc1p1OEJLa0JQdyJ9.eyJhdWQiOiJhcGk6Ly83NzMxZWM3NC1mMWNhLTQxZDUtYTBlNS04MWY2ODM3YWMyNDciLCJpc3MiOiJodHRwczovL3N0cy53aW5kb3dzLm5ldC83MDExNTk1NC0wY2NkLTQ1ZjAtODdiZC0wM2IyYTM1ODc1NjkvIiwiaWF0IjoxNzM2NDc1NTk1LCJuYmYiOjE3MzY0NzU1OTUsImV4cCI6MTczNjQ4MDUyMSwiYWNyIjoiMSIsImFpbyI6IkFWUUFxLzhZQUFBQVhaaHZKNFlNcEZiVWZ3V2RQQnVIcHFHMHd6bUQ5WGNRbjBZbUZ6eUVKL2hMZ1JqOXBNcHduWGZwc1A0dlAzcmVTVWFvdXNQMDNEbUtWb2VmLzRwSENTQ0VueTRiM0p5SXNQOXU1Zzk5Z3BVPSIsImFtciI6WyJwd2QiLCJyc2EiLCJtZmEiXSwiYXBwaWQiOiI3NzMxZWM3NC1mMWNhLTQxZDUtYTBlNS04MWY2ODM3YWMyNDciLCJhcHBpZGFjciI6IjAiLCJkZXZpY2VpZCI6IjMyOTk1Y2U1LTM0NjMtNGE4Yy1iM2MyLWFkNjRlZmY0ZjhkNyIsImZhbWlseV9uYW1lIjoiWWFuZyIsImdpdmVuX25hbWUiOiJKZW5ueSIsImlwYWRkciI6IjIwNS4xNzguMzAuNDUiLCJuYW1lIjoiWWFuZywgSmVubnkiLCJvaWQiOiI1YTVlNDNjNi0wMWEzLTQ4MjItOGE0My0yY2E2OThlYjViNTYiLCJvbnByZW1fc2lkIjoiUy0xLTUtMjEtMzYzODA4OTg2OC0zMDgxMTc1MTE1LTMxMzI5NzcyOS0zMjE1MzMiLCJyaCI6IjEuQVJzQVZGa1JjTTBNOEVXSHZRT3lvMWgxYVhUc01YZks4ZFZCb09XQjlvTjZ3a2NiQU9rYkFBLiIsInNjcCI6IkFwaS5SZWFkIiwic3ViIjoiRTFWeklZblc2RWI0aE5ZU0xZNE9LUTVBbXQ4SVBUaVRPWXZCSHZaNzYxZyIsInRpZCI6IjcwMTE1OTU0LTBjY2QtNDVmMC04N2JkLTAzYjJhMzU4NzU2OSIsInVuaXF1ZV9uYW1lIjoiMTA0Mjg4QGdsb2JhbC51bC5jb20iLCJ1cG4iOiIxMDQyODhAZ2xvYmFsLnVsLmNvbSIsInV0aSI6IjZFMUxUOTlyOGtxUDJnSmRRck1mQUEiLCJ2ZXIiOiIxLjAifQ.Ap9oGLSD8E-nz7i4g5Ksvfx9sr-SOgCUm4Nf6Wjc-a3j_TmBX8lvQHeu7vj_miIjip1jJU-4RJqnriEZrmUkuTTC1tL17KTHCh1ifiSHbwlrYasLr-BXzKhhUj3mugh_LCOK80KVBT7NNp63BTd0EyWwLhDWT492ns-kJ0afW0pCwBg6hBx63Uz9i6sW6WxqSNAD_xSVGkCR5x8wnwrAwAneTpUsnn2k3bCA0qk0Fis1wWOsakPF8Z6fo2yBRQAlcsizalj59FylAeJMSD56fKPIxEZw3qXQlrlR8huh-67_VwJfeSnJ7Qfx_jsZYavZ0W3C_3O0fglx1YhHK7DR4Q'
    user_id = '5a5e43c6-01a3-4822-8a43-2ca698eb5b56'         # Replace with your actual user ID
    
    # Create an instance of the LLMRequest class
    llm_request = LLMRequest(auth_token, user_id)
    
    # Call the method to compare regulatory information
    llm_request.compare_regulatory_info()
